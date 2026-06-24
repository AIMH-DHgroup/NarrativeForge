from __future__ import annotations

import html
from pathlib import Path
from typing import Any

import pandas as pd

from .utils import compact_float, percentage


def _row(table: pd.DataFrame, sort_col: str | None = None, ascending: bool = False) -> pd.Series:
    if table.empty:
        return pd.Series(dtype=object)
    if sort_col and sort_col in table.columns:
        return table.sort_values(sort_col, ascending=ascending).iloc[0]
    return table.iloc[0]


def _fmt_config(row: pd.Series) -> str:
    if row.empty:
        return "not available"
    parts = []
    if "model" in row and pd.notna(row.get("model")):
        parts.append(str(row.get("model")))
    if "input_strategy" in row and pd.notna(row.get("input_strategy")):
        parts.append(str(row.get("input_strategy")))
    if "prompt_strategy" in row and pd.notna(row.get("prompt_strategy")):
        parts.append(str(row.get("prompt_strategy")))
    label = " + ".join(parts) if parts else "configuration"
    nrs = compact_float(row.get("mean_NRS"), 3)
    runtime = compact_float(row.get("mean_runtime_seconds"), 3)
    param_value = row.get("parameters_b", row.get("parameter_median_b", None))
    params = compact_float(param_value, 3)
    param_text = "NA" if params == "NA" else f"{params}B"
    fail = percentage(row.get("failure_rate"), 1)
    return f"{label} (mean NRS={nrs}, runtime={runtime}s, parameters={param_text}, failure rate={fail})"


def _df_text(df: pd.DataFrame, rows: int = 12) -> str:
    if df.empty:
        return "No data available."
    preview = df.head(rows).copy()
    return preview.to_string(index=False, max_cols=12)


def _df_html(df: pd.DataFrame, rows: int = 20) -> str:
    if df.empty:
        return "<p>No data available.</p>"
    preview = df.head(rows).copy()
    return preview.to_html(index=False, border=0, classes="dataframe", float_format=lambda x: f"{x:.3f}")


def _best_tradeoff_row(tables: dict[str, pd.DataFrame], criterion: str) -> pd.Series:
    table = tables.get("best_tradeoff_configurations", pd.DataFrame())
    if table.empty or "criterion" not in table.columns:
        return pd.Series(dtype=object)
    subset = table[table["criterion"] == criterion]
    return subset.iloc[0] if not subset.empty else pd.Series(dtype=object)


def _generate_findings(df: pd.DataFrame, tables: dict[str, pd.DataFrame], summary: dict[str, Any], options) -> dict[str, str]:
    input_best = _row(tables.get("input_strategy_summary", pd.DataFrame()), "mean_NRS", False)
    prompt_best = _row(tables.get("prompt_strategy_summary", pd.DataFrame()), "mean_NRS", False)
    input_prompt_best = _row(tables.get("input_prompt_summary", pd.DataFrame()), "mean_NRS", False)
    model_best = _row(tables.get("model_summary", pd.DataFrame()), "mean_NRS", False)
    raw_best = _best_tradeoff_row(tables, "best raw quality")
    balanced_best = _best_tradeoff_row(tables, "best balanced score")
    reliable_balanced_best = _best_tradeoff_row(tables, "best reliable balanced score")
    qps_best = _best_tradeoff_row(tables, "best quality per second")
    qpb_best = _best_tradeoff_row(tables, "best quality per billion parameters")

    quality_loss = tables.get("quality_loss_by_parameter_threshold", pd.DataFrame())
    small_model_sentence = "No model-size threshold table was available."
    if not quality_loss.empty:
        rows = []
        for _, row in quality_loss.iterrows():
            rows.append(
                f"{row.get('threshold')}: {row.get('best_model')} + {row.get('best_input_strategy')} + {row.get('best_prompt_strategy')} "
                f"lost {compact_float(row.get('NRS_loss_vs_global_best'), 3)} NRS points"
            )
        small_model_sentence = "; ".join(rows) + "."

    return {
        "rq1": f"The highest aggregate input strategy is {_fmt_config(input_best)}.",
        "rq2": f"The highest aggregate prompt strategy is {_fmt_config(prompt_best)}.",
        "rq3": f"The strongest aggregate input-by-prompt cell is {_fmt_config(input_prompt_best)}.",
        "rq4": f"The highest overall model is {_fmt_config(model_best)}.",
        "rq7": small_model_sentence,
        "raw_best": _fmt_config(raw_best),
        "balanced_best": _fmt_config(balanced_best),
        "reliable_balanced_best": _fmt_config(reliable_balanced_best),
        "qps_best": _fmt_config(qps_best),
        "qpb_best": _fmt_config(qpb_best),
    }


def _scientific_interpretation(tables: dict[str, pd.DataFrame]) -> list[str]:
    matrix = tables.get("model_strategy_matrix", pd.DataFrame())
    family = tables.get("model_family_summary", pd.DataFrame())
    top5 = tables.get("top5_configurations", pd.DataFrame())
    lines: list[str] = []
    if not matrix.empty:
        best = matrix.iloc[0]
        lines.append(
            f"The strongest model-level result is {best.get('model')} with best_NRS={compact_float(best.get('best_NRS'), 3)} under {best.get('best_strategy')} input."
        )
        available = [c for c in ["auto_NRS", "brief_NRS", "rag_NRS", "full_NRS"] if c in matrix.columns]
        if available:
            winners = matrix[["model", *available]].copy()
            lines.append("Strategy effects are model-dependent; the model-strategy matrix should be used to check whether RAG or full context benefits all models equally.")
    if not family.empty:
        fam = family.iloc[0]
        lines.append(
            f"At family level, {fam.get('family')} reaches the highest observed max_NRS={compact_float(fam.get('max_NRS'), 3)} across {int(fam.get('n_models', 0) or 0)} model(s)."
        )
        lines.append("Scaling curves reveal whether larger models consistently improve NRS or whether quality plateaus within each family, indicating diminishing returns.")
    if not top5.empty:
        top = top5.iloc[0]
        lines.append(
            f"The best configuration is {top.get('model')} + {top.get('input_strategy')} + {top.get('prompt_strategy')} with mean_NRS={compact_float(top.get('mean_NRS'), 3)} and runtime={compact_float(top.get('runtime_seconds'), 3)}s."
        )
        fastest = top5.sort_values("runtime_seconds", ascending=True).iloc[0]
        if fastest.get("rank") != top.get("rank"):
            lines.append(
                f"The highest-NRS configuration is not the fastest among the top five; {fastest.get('model')} + {fastest.get('input_strategy')} + {fastest.get('prompt_strategy')} is the fastest top-five option."
            )
        else:
            lines.append("The best-NRS configuration is also the fastest within the top-five subset, making it unusually efficient.")
    lines.append("Smaller models remain competitive when they appear near the top of the strategy matrix or Pareto frontier despite lower parameter counts.")
    return lines


COVERAGE_LIMITATION_TEXT = (
    "BERTScore F1 and sentence-transformer semantic similarity were used as automatic semantic-alignment indicators between the generated narrative and the source document. "
    "They estimate whether the generated text preserves the meaning of the source at a global semantic level. However, because the comparison is performed between the complete generated narrative and the complete source document, these metrics may overestimate quality when the source is much longer than the generated narrative: a text can remain topically similar while omitting important facts, causal relations, or local details. "
    "Consequently, these measures do not replace expert factual validation, claim-level checking, or human evaluation of narrative faithfulness."
)


def _coverage_interpretation(tables: dict[str, pd.DataFrame]) -> list[str]:
    input_cov = tables.get("coverage_by_input_strategy", pd.DataFrame())
    model_cov = tables.get("coverage_by_model", pd.DataFrame())
    suspicious = tables.get("high_alignment_low_coverage_cases", pd.DataFrame())
    lines = [COVERAGE_LIMITATION_TEXT]
    if not input_cov.empty and "mean_source_coverage_075" in input_cov.columns:
        best = input_cov.sort_values("mean_source_coverage_075", ascending=False).iloc[0]
        lines.append(f"The highest mean source coverage@0.75 is observed for input strategy {best.get('input_strategy')} ({compact_float(best.get('mean_source_coverage_075'), 3)}).")
    if not model_cov.empty and "mean_omission_risk_075" in model_cov.columns:
        risk = model_cov.sort_values("mean_omission_risk_075", ascending=False).iloc[0]
        lines.append(f"The largest mean omission risk@0.75 is observed for {risk.get('model')} ({compact_float(risk.get('mean_omission_risk_075'), 3)}), indicating possible global similarity without broad source coverage.")
    if not suspicious.empty:
        lines.append(f"There are {len(suspicious)} high-alignment/low-coverage cases with BERTScore F1 >= 0.85 and source coverage@0.75 <= 0.50.")
    return lines


def _case_study_interpretation(tables: dict[str, pd.DataFrame]) -> list[str]:
    summary = tables.get("case_study_summary", pd.DataFrame())
    difficulty = tables.get("case_difficulty_ranking", pd.DataFrame())
    stability = tables.get("case_stability_summary", pd.DataFrame())
    input_delta = tables.get("case_input_strategy_delta", pd.DataFrame())
    prompt_delta = tables.get("case_prompt_strategy_delta", pd.DataFrame())
    size_loss = tables.get("case_model_size_loss", pd.DataFrame())
    best = tables.get("best_configuration_by_case", pd.DataFrame())
    if summary.empty:
        return ["Case-study analysis was unavailable because no case-study identifier column was found."]
    lines = [
        f"The benchmark contains {len(summary)} case studies. Case-level aggregation reveals whether benchmark conclusions are driven by a small subset of cases or whether observed strategy effects are consistent across scenarios. A case with low mean NRS and high variance should be treated as intrinsically difficult or underspecified, while a case with high input-strategy range is particularly sensitive to retrieval or context length. Conversely, a case where small models remain close to the case-level optimum suggests that the task may depend more on source coverage and prompt structure than on model scale."
    ]
    easiest = summary.sort_values("mean_NRS", ascending=False).iloc[0]
    hardest = summary.sort_values("mean_NRS", ascending=True).iloc[0]
    lines.append(f"The easiest case by mean NRS is {easiest.get('case_study')} ({compact_float(easiest.get('mean_NRS'), 3)}), while the hardest is {hardest.get('case_study')} ({compact_float(hardest.get('mean_NRS'), 3)}).")
    fail = summary.sort_values("failure_rate", ascending=False).iloc[0]
    lines.append(f"The highest failure rate occurs for {fail.get('case_study')} ({percentage(fail.get('failure_rate'), 1)}).")
    if not difficulty.empty:
        hard = difficulty.iloc[0]
        easy = difficulty.sort_values("difficulty_score", ascending=True).iloc[0]
        lines.append(f"The difficulty score ranks {hard.get('case_study')} as hardest and {easy.get('case_study')} as easiest after combining inverse NRS, failure rate, and NRS variability.")
    if not stability.empty:
        most = stability.sort_values("stability_score", ascending=False).iloc[0]
        least = stability.sort_values("stability_score", ascending=True).iloc[0]
        lines.append(f"The most stable case is {most.get('case_study')}; the least stable is {least.get('case_study')}, based on normalized NRS standard deviation.")
    if not input_delta.empty:
        row = input_delta.sort_values("input_strategy_NRS_range", ascending=False).iloc[0]
        lines.append(f"The case most sensitive to input strategy is {row.get('case_study')} with an NRS range of {compact_float(row.get('input_strategy_NRS_range'), 3)}.")
    if not prompt_delta.empty:
        row = prompt_delta.sort_values("prompt_strategy_NRS_range", ascending=False).iloc[0]
        lines.append(f"The case most sensitive to prompt strategy is {row.get('case_study')} with an NRS range of {compact_float(row.get('prompt_strategy_NRS_range'), 3)}.")
    if not size_loss.empty and "NRS_loss_vs_case_best" in size_loss.columns:
        constrained = size_loss[size_loss["threshold"] != "all"].dropna(subset=["NRS_loss_vs_case_best"])
        if not constrained.empty:
            row = constrained.sort_values("NRS_loss_vs_case_best", ascending=False).iloc[0]
            competitive = constrained.sort_values("NRS_loss_vs_case_best", ascending=True).iloc[0]
            lines.append(f"The largest small-model quality loss appears for {row.get('case_study')} under {row.get('threshold')} ({compact_float(row.get('NRS_loss_vs_case_best'), 3)} NRS points). Small models remain most competitive for {competitive.get('case_study')} under {competitive.get('threshold')}.")
    if not best.empty:
        lines.append("The best-configuration-by-case table should be used to check whether the global best model also dominates individual case studies or whether different cases require different models and strategies.")
    return lines


def _append_unlisted_figures_markdown(lines: list[str], summary: dict[str, Any], listed: set[str]) -> None:
    for fig_key, fig_value in summary.get("figures", {}).items():
        if fig_key in listed:
            continue
        fig_path = Path(fig_value)
        rel = f"figures/{fig_path.parent.name}/{fig_path.name}" if fig_path.parent.name != "figures" else f"figures/{fig_path.name}"
        lines.append(f"### {fig_key.replace('_', ' ').title()}\n")
        lines.append(f"![{fig_key}]({rel})\n")


def _append_unlisted_figures_html(html_parts: list[str], summary: dict[str, Any], listed: set[str]) -> None:
    for fig_key, fig_value in summary.get("figures", {}).items():
        if fig_key in listed:
            continue
        fig_path = Path(fig_value)
        rel = f"figures/{fig_path.parent.name}/{fig_path.name}" if fig_path.parent.name != "figures" else f"figures/{fig_path.name}"
        html_parts.append("<div class='card'>")
        html_parts.append(f"<h3>{html.escape(fig_key.replace('_', ' ').title())}</h3>")
        html_parts.append(f"<img src='{html.escape(rel)}' alt='{html.escape(fig_key)}'>")
        html_parts.append("</div>")


def write_markdown_report(path: str | Path, df: pd.DataFrame, tables: dict[str, pd.DataFrame], summary: dict[str, Any], options) -> Path:
    path = Path(path)
    findings = _generate_findings(df, tables, summary, options)
    desc = tables.get("descriptive_statistics", pd.DataFrame())
    corr = tables.get("parameter_correlation_summary", pd.DataFrame())

    lines: list[str] = []
    lines.append(f"# {options.title}\n")
    lines.append("## Scope and data handling\n")
    lines.append(
        "This report was generated by NarrativeForge NRS Analyzer. The software loads all `nrs_runs.csv` files from a ZIP archive or extracted experiment directory, derives `input_strategy` from `benchmark_all_*` parent folders when available, concatenates all runs, and uses `NRS` as the primary quality metric. `CSV_NRS` is not used for the main conclusions. Failed runs are retained in all-attempt summaries and are separately represented through the `failure_rate` field.\n"
    )
    lines.append(f"Hardware/runtime note: {options.hardware_note}\n")

    lines.append("## Descriptive statistics\n")
    lines.append("```text\n" + _df_text(desc, rows=5) + "\n```\n")

    lines.append("## Research-question answers\n")
    lines.append(f"**RQ1 — best input strategy.** {findings['rq1']}\n")
    lines.append(f"**RQ2 — best prompt strategy.** {findings['rq2']}\n")
    lines.append(f"**RQ3 — best input x prompt pair.** {findings['rq3']}\n")
    lines.append(f"**RQ4 — best model.** {findings['rq4']}\n")
    lines.append(
        "**RQ5 — model size and NRS.** Parameter-count effects should be interpreted from the configuration and model-size tables. A positive correlation indicates that larger models tend to improve NRS, whereas a weak or plateauing relationship indicates that model family, prompt design, and input strategy matter as much as raw size.\n"
    )
    lines.append(
        "**RQ6 — model size and runtime.** Runtime is expected to increase with parameter count, but the slope depends on input length and prompt strategy. The runtime-vs-parameters figures should therefore be interpreted at the configuration level, not only at the model level.\n"
    )
    lines.append(f"**RQ7 — quality loss under size constraints.** {findings['rq7']}\n")
    lines.append(
        f"**RQ8 — quality-efficiency tradeoff.** Best raw-quality configuration: {findings['raw_best']}. Best balanced configuration: {findings['balanced_best']}. Best reliable balanced configuration: {findings['reliable_balanced_best']}. Best quality-per-second configuration: {findings['qps_best']}. Best quality-per-billion-parameters configuration: {findings['qpb_best']}.\n"
    )
    lines.append(
        "**RQ9 — deployment on normal computers.** A normal-computer deployment should generally prefer Pareto-optimal, reliable small-to-mid-size configurations rather than the global raw-quality maximum. The raw-quality maximum may be useful as an upper-bound reference, but it should be retested on the actual target machine.\n"
    )

    lines.append("## Strategy analysis\n")
    lines.append("### Mean NRS and runtime by input strategy\n")
    lines.append("```text\n" + _df_text(tables.get("input_strategy_summary", pd.DataFrame())) + "\n```\n")
    lines.append("### Mean NRS and runtime by prompt strategy\n")
    lines.append("```text\n" + _df_text(tables.get("prompt_strategy_summary", pd.DataFrame())) + "\n```\n")
    lines.append("### Input strategy x prompt strategy\n")
    lines.append("```text\n" + _df_text(tables.get("input_prompt_summary", pd.DataFrame()), rows=20) + "\n```\n")

    lines.append("## Model and model-size analysis\n")
    lines.append("### Mean NRS by model\n")
    lines.append("```text\n" + _df_text(tables.get("model_summary", pd.DataFrame()), rows=20) + "\n```\n")
    lines.append("### Parameter correlations\n")
    lines.append("```text\n" + _df_text(corr, rows=5) + "\n```\n")
    lines.append("### Model-size bins\n")
    lines.append("```text\n" + _df_text(tables.get("size_bin_summary", pd.DataFrame()), rows=10) + "\n```\n")

    lines.append("## Model Family Analysis\n")
    for text in _scientific_interpretation(tables):
        lines.append(f"- {text}\n")
    lines.append("### Global model strategy matrix\n")
    lines.append("```text\n" + _df_text(tables.get("model_strategy_matrix", pd.DataFrame()), rows=30) + "\n```\n")
    lines.append("### Prompt-specific model matrix\n")
    lines.append("```text\n" + _df_text(tables.get("model_strategy_prompt_matrix", pd.DataFrame()), rows=30) + "\n```\n")
    lines.append("### Model family summary\n")
    lines.append("```text\n" + _df_text(tables.get("model_family_summary", pd.DataFrame()), rows=20) + "\n```\n")

    lines.append("## Scaling Laws\n")
    lines.append("Family scaling figures show NRS, runtime, and quality-efficiency curves within each detected model family. Use these plots to identify whether extra parameters produce consistent gains or diminishing returns.\n")

    lines.append("## Top-5 Configuration Analysis\n")
    lines.append("```text\n" + _df_text(tables.get("top5_configurations", pd.DataFrame()), rows=5) + "\n```\n")

    lines.append("## Top-5 Model Analysis\n")
    lines.append("```text\n" + _df_text(tables.get("top_models_strategy_matrix", pd.DataFrame()), rows=5) + "\n```\n")

    lines.append("## Efficiency and Pareto analysis\n")
    lines.append(
        "The balanced score is a configurable deployment-oriented index: `0.50 * normalized_NRS + 0.20 * normalized_speed + 0.15 * normalized_size_efficiency + 0.10 * normalized_stability + 0.05 * normalized_reliability` by default. The weights can be changed from the CLI or dashboard.\n"
    )
    lines.append("### Best tradeoff configurations\n")
    lines.append("```text\n" + _df_text(tables.get("best_tradeoff_configurations", pd.DataFrame()), rows=10) + "\n```\n")
    lines.append("### Quality loss by parameter threshold\n")
    lines.append("```text\n" + _df_text(tables.get("quality_loss_by_parameter_threshold", pd.DataFrame()), rows=10) + "\n```\n")
    lines.append("### Pareto-optimal configurations\n")
    lines.append("```text\n" + _df_text(tables.get("pareto_optimal_configurations", pd.DataFrame()), rows=30) + "\n```\n")

    lines.append("## Coverage Diagnostics and Omission Risk\n")
    for text in _coverage_interpretation(tables):
        lines.append(f"{text}\n")
    for title, key in [
        ("Coverage by input strategy", "coverage_by_input_strategy"),
        ("Coverage by prompt strategy", "coverage_by_prompt_strategy"),
        ("Coverage by model", "coverage_by_model"),
        ("Coverage by model family", "coverage_by_model_family"),
        ("High-alignment low-coverage cases", "high_alignment_low_coverage_cases"),
    ]:
        lines.append(f"### {title}\n")
        lines.append("```text\n" + _df_text(tables.get(key, pd.DataFrame()), rows=20) + "\n```\n")

    lines.append("## Case Study Analysis\n")
    for text in _case_study_interpretation(tables):
        lines.append(f"{text}\n")
    for title, key in [
        ("Case-study summary", "case_study_summary"),
        ("Case difficulty ranking", "case_difficulty_ranking"),
        ("Best configuration by case", "best_configuration_by_case"),
        ("Case input-strategy deltas", "case_input_strategy_delta"),
        ("Case prompt-strategy deltas", "case_prompt_strategy_delta"),
        ("Case model-size loss", "case_model_size_loss"),
    ]:
        lines.append(f"### {title}\n")
        lines.append("```text\n" + _df_text(tables.get(key, pd.DataFrame()), rows=30) + "\n```\n")

    lines.append("## Figures\n")
    listed_figures: set[str] = set()
    for item in summary.get("figure_interpretations", []):
        fig_key = item["figure"]
        listed_figures.add(fig_key)
        fig_path = Path(summary.get("figures", {}).get(fig_key, ""))
        rel = f"figures/{fig_path.name}" if fig_path.name else ""
        lines.append(f"### {item['title']}\n")
        if rel:
            lines.append(f"![{item['title']}]({rel})\n")
        lines.append(item.get("interpretation", "") + "\n")
    _append_unlisted_figures_markdown(lines, summary, listed_figures)

    lines.append("## Critical discussion\n")
    lines.append(
        "The benchmark is strong because it evaluates multiple input strategies, prompt strategies, models, runtimes, failures, and repeated configurations in a single experimental design. This makes it possible to distinguish raw quality from deployment feasibility. However, NRS remains an aggregate metric and should not be treated as a complete substitute for human evaluation. Aggregate metrics can obscure failure modes such as factual inconsistency, stylistic drift, verbosity, missing constraints, or brittle behavior on rare cases.\n"
    )
    lines.append(
        "BERTScore and semantic-similarity components can reward lexical or semantic proximity without fully capturing narrative coherence, causal structure, pacing, or reader-perceived quality. They are valuable as scalable signals but should be complemented with targeted qualitative review. Runtime is hardware-dependent; measurements collected on a high-performance GPU workstation should be interpreted comparatively within this experiment, not as expected runtimes on laptops or commodity desktops.\n"
    )
    lines.append(
        "Threats to validity include benchmark-case representativeness, model-family confounding, prompt-template confounding, parameter-count inference errors for nonstandard model names, and runtime variability from local system load. Future work should retest the Pareto-optimal and balanced configurations on commodity hardware, add human preference evaluation, test more narrative genres, evaluate longer-context failure modes, and analyze confidence intervals over independent benchmark replications.\n"
    )
    lines.append(
        "Coverage diagnostics reduce but do not eliminate factuality threats. Future work should add claim-level factuality evaluation, source-span attribution, coverage-based scoring, human expert review, section-level or paragraph-level semantic comparison, and entity and event consistency checks.\n"
    )

    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_html_report(path: str | Path, df: pd.DataFrame, tables: dict[str, pd.DataFrame], summary: dict[str, Any], options) -> Path:
    path = Path(path)
    findings = _generate_findings(df, tables, summary, options)
    css = """
    body { font-family: system-ui, -apple-system, BlinkMacSystemFont, Segoe UI, sans-serif; max-width: 1180px; margin: 32px auto; line-height: 1.55; color: #1f2933; }
    h1, h2, h3 { color: #102a43; }
    .card { border: 1px solid #d9e2ec; border-radius: 12px; padding: 16px 20px; margin: 18px 0; background: #f8fafc; }
    .grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 14px; }
    .metric { border: 1px solid #e2e8f0; border-radius: 10px; padding: 12px; background: white; }
    .metric strong { display: block; font-size: 0.9rem; color: #52606d; }
    img { max-width: 100%; border: 1px solid #d9e2ec; border-radius: 8px; margin-top: 8px; }
    table.dataframe { border-collapse: collapse; width: 100%; font-size: 0.85rem; overflow-x: auto; display: block; }
    table.dataframe th, table.dataframe td { border: 1px solid #d9e2ec; padding: 6px 8px; text-align: right; white-space: nowrap; }
    table.dataframe th { background: #edf2f7; color: #243b53; }
    table.dataframe td:first-child, table.dataframe th:first-child { text-align: left; }
    .note { color: #52606d; font-size: 0.95rem; }
    """

    desc = tables.get("descriptive_statistics", pd.DataFrame())
    html_parts = ["<!doctype html><html><head><meta charset='utf-8'>", f"<title>{html.escape(options.title)}</title>", f"<style>{css}</style>", "</head><body>"]
    html_parts.append(f"<h1>{html.escape(options.title)}</h1>")
    html_parts.append(
        "<p>This report was generated by NarrativeForge NRS Analyzer. It uses <strong>NRS</strong> as the primary quality metric, <strong>runtime_seconds</strong> as the execution-time metric, and treats failed runs separately through failure-rate summaries.</p>"
    )
    html_parts.append(f"<p class='note'>{html.escape(options.hardware_note)}</p>")

    html_parts.append("<h2>Descriptive statistics</h2>")
    html_parts.append(_df_html(desc, rows=5))

    html_parts.append("<h2>Research-question answers</h2><div class='grid'>")
    cards = [
        ("RQ1: Input strategy", findings["rq1"]),
        ("RQ2: Prompt strategy", findings["rq2"]),
        ("RQ3: Input x prompt", findings["rq3"]),
        ("RQ4: Model", findings["rq4"]),
        ("RQ7: Size-constrained loss", findings["rq7"]),
        ("RQ8: Balanced deployment", findings["balanced_best"] + " Reliable balanced: " + findings["reliable_balanced_best"]),
    ]
    for title, text in cards:
        html_parts.append(f"<div class='metric'><strong>{html.escape(title)}</strong>{html.escape(text)}</div>")
    html_parts.append("</div>")

    html_parts.append("<h2>Core tables</h2>")
    for title, key, rows in [
        ("Input strategy summary", "input_strategy_summary", 20),
        ("Prompt strategy summary", "prompt_strategy_summary", 20),
        ("Input x prompt summary", "input_prompt_summary", 30),
        ("Model summary", "model_summary", 30),
        ("Global model strategy matrix", "model_strategy_matrix", 50),
        ("Prompt-specific model matrix", "model_strategy_prompt_matrix", 50),
        ("Model family summary", "model_family_summary", 30),
        ("Top-5 configurations", "top5_configurations", 5),
        ("Top models strategy heatmap table", "top_models_strategy_matrix", 5),
        ("Coverage by input strategy", "coverage_by_input_strategy", 20),
        ("Coverage by prompt strategy", "coverage_by_prompt_strategy", 20),
        ("Coverage by model", "coverage_by_model", 30),
        ("High-alignment low-coverage cases", "high_alignment_low_coverage_cases", 30),
        ("Best tradeoff configurations", "best_tradeoff_configurations", 10),
        ("Quality loss by parameter threshold", "quality_loss_by_parameter_threshold", 10),
        ("Pareto-optimal configurations", "pareto_optimal_configurations", 40),
    ]:
        html_parts.append(f"<h3>{html.escape(title)}</h3>")
        html_parts.append(_df_html(tables.get(key, pd.DataFrame()), rows=rows))

    html_parts.append("<h2>Model Family Analysis</h2>")
    for text in _scientific_interpretation(tables):
        html_parts.append(f"<p>{html.escape(text)}</p>")
    html_parts.append("<h2>Scaling Laws</h2>")
    html_parts.append("<p>Family scaling figures show NRS, runtime, and quality-efficiency curves within each detected model family, making diminishing returns and efficient model sizes visible.</p>")
    html_parts.append("<h2>Top-5 Configuration Analysis</h2>")
    html_parts.append(_df_html(tables.get("top5_configurations", pd.DataFrame()), rows=5))
    html_parts.append("<h2>Top-5 Model Analysis</h2>")
    html_parts.append(_df_html(tables.get("top_models_strategy_matrix", pd.DataFrame()), rows=5))

    html_parts.append("<h2>Coverage Diagnostics and Omission Risk</h2>")
    for text in _coverage_interpretation(tables):
        html_parts.append(f"<p>{html.escape(text)}</p>")
    html_parts.append(_df_html(tables.get("coverage_by_input_strategy", pd.DataFrame()), rows=20))
    html_parts.append(_df_html(tables.get("high_alignment_low_coverage_cases", pd.DataFrame()), rows=30))

    html_parts.append("<h2>Case Study Analysis</h2>")
    for text in _case_study_interpretation(tables):
        html_parts.append(f"<p>{html.escape(text)}</p>")
    for title, key, rows in [
        ("Case-study summary", "case_study_summary", 30),
        ("Case difficulty ranking", "case_difficulty_ranking", 30),
        ("Best configuration by case", "best_configuration_by_case", 30),
        ("Case input-strategy deltas", "case_input_strategy_delta", 30),
        ("Case prompt-strategy deltas", "case_prompt_strategy_delta", 30),
        ("Case model-size loss", "case_model_size_loss", 60),
    ]:
        html_parts.append(f"<h3>{html.escape(title)}</h3>")
        html_parts.append(_df_html(tables.get(key, pd.DataFrame()), rows=rows))

    html_parts.append("<h2>Figures and interpretations</h2>")
    listed_figures: set[str] = set()
    for item in summary.get("figure_interpretations", []):
        fig_key = item["figure"]
        listed_figures.add(fig_key)
        fig_path = Path(summary.get("figures", {}).get(fig_key, ""))
        rel = f"figures/{fig_path.name}" if fig_path.name else ""
        html_parts.append("<div class='card'>")
        html_parts.append(f"<h3>{html.escape(item.get('title', fig_key))}</h3>")
        if rel:
            html_parts.append(f"<img src='{html.escape(rel)}' alt='{html.escape(item.get('title', fig_key))}'>")
        html_parts.append(f"<p>{html.escape(item.get('interpretation', ''))}</p>")
        html_parts.append("</div>")
    _append_unlisted_figures_html(html_parts, summary, listed_figures)

    html_parts.append("<h2>Critical discussion</h2>")
    html_parts.append(
        "<p>The benchmark design supports comparative evaluation across input strategies, prompt strategies, models, runtimes, failures, and stability. The main limitation is that NRS is an aggregate automatic metric. It is useful for scalable comparison but should be complemented with targeted human evaluation, especially for narrative coherence, factual consistency, style, pacing, and rare-case failures.</p>"
    )
    html_parts.append(
        "<p>BERTScore and semantic similarity signals can capture semantic overlap but do not fully measure reader-perceived quality or discourse-level structure. Runtime is also hardware-dependent, so high-performance GPU workstation measurements should be treated as relative rankings unless deployment hardware is equivalent.</p>"
    )
    html_parts.append(
        "<p>Recommended next steps are to retest Pareto-optimal and balanced configurations on commodity hardware, repeat the benchmark with additional cases, add human preference review, and test longer-context and domain-shifted narratives.</p>"
    )
    html_parts.append(
        "<p>Future work should also include claim-level factuality evaluation, source-span attribution, coverage-based scoring, section-level or paragraph-level semantic comparison, and entity and event consistency checks.</p>"
    )
    html_parts.append("</body></html>")
    path.write_text("\n".join(html_parts), encoding="utf-8")
    return path


def write_latex_report(path: str | Path, df: pd.DataFrame, tables: dict[str, pd.DataFrame], summary: dict[str, Any], options) -> Path:
    path = Path(path)
    lines = [r"\documentclass{article}", r"\usepackage{booktabs}", r"\usepackage{graphicx}", r"\usepackage[margin=1in]{geometry}", r"\begin{document}", f"\\title{{{options.title}}}", r"\maketitle"]
    lines.append(r"\section{Model Family Analysis}")
    for text in _scientific_interpretation(tables):
        lines.append(text.replace("_", r"\_"))
        lines.append("\n")
    for title, key in [("Global Model Strategy Matrix", "model_strategy_matrix"), ("Family Summary", "model_family_summary"), ("Top-5 Configurations", "top5_configurations")]:
        table = tables.get(key, pd.DataFrame())
        if table.empty:
            continue
        lines.append(f"\\section{{{title}}}")
        lines.append(table.head(20).to_latex(index=False, float_format="%.3f", escape=True))
    lines.append(r"\section{Coverage Diagnostics and Omission Risk}")
    for text in _coverage_interpretation(tables):
        lines.append(text.replace("_", r"\_"))
        lines.append("\n")
    cov = tables.get("coverage_by_input_strategy", pd.DataFrame())
    if not cov.empty:
        lines.append(cov.head(20).to_latex(index=False, float_format="%.3f", escape=True))
    lines.append(r"\section{Case Study Analysis}")
    for text in _case_study_interpretation(tables):
        lines.append(text.replace("_", r"\_"))
        lines.append("\n")
    for key in ["case_study_summary", "case_difficulty_ranking", "best_configuration_by_case", "case_input_strategy_delta", "case_prompt_strategy_delta", "case_model_size_loss"]:
        table = tables.get(key, pd.DataFrame())
        if not table.empty:
            lines.append(table.head(20).to_latex(index=False, float_format="%.3f", escape=True))
    lines.append(r"\end{document}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_docx_report(path: str | Path, df: pd.DataFrame, tables: dict[str, pd.DataFrame], summary: dict[str, Any], options) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None
    path = Path(path)
    document = Document()
    document.add_heading(options.title, 0)
    document.add_heading("Model Family Analysis", level=1)
    for text in _scientific_interpretation(tables):
        document.add_paragraph(text)
    for title, key in [("Global model strategy matrix", "model_strategy_matrix"), ("Model family summary", "model_family_summary"), ("Top-5 configurations", "top5_configurations")]:
        table_df = tables.get(key, pd.DataFrame()).head(10)
        document.add_heading(title, level=1)
        if table_df.empty:
            document.add_paragraph("No data available.")
            continue
        table = document.add_table(rows=1, cols=len(table_df.columns))
        for i, col in enumerate(table_df.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in table_df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = compact_float(value, 3) if isinstance(value, float) else str(value)
    document.add_heading("Coverage Diagnostics and Omission Risk", level=1)
    for text in _coverage_interpretation(tables):
        document.add_paragraph(text)
    document.add_heading("Case Study Analysis", level=1)
    for text in _case_study_interpretation(tables):
        document.add_paragraph(text)
    document.save(path)
    return path
