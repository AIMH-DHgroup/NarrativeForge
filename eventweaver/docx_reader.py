from __future__ import annotations

from pathlib import Path


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Install it with: pip install python-docx") from exc

    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_document(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() in {".txt", ".md", ".text"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported document type: {path}")
