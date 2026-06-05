from pathlib import Path
from docx import Document
import re
import sys


def count_docx(path: Path) -> tuple[int, int]:
    doc = Document(path)

    parts = []

    # Paragraph text
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)

    # Table text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    text = "\n".join(parts)

    words = re.findall(r"\b\w+\b", text)
    chars = len(text)

    return len(words), chars


def main(folder: str) -> None:
    folder_path = Path(folder)

    if not folder_path.is_dir():
        print(f"Error: not a folder: {folder}")
        sys.exit(1)

    print("document\twords\tchars")

    for docx_file in sorted(folder_path.glob("*.docx")):
        if docx_file.name.startswith("~$"):
            continue  # skip temporary Word files

        words, chars = count_docx(docx_file)
        print(f"{docx_file.name}\t{words}\t{chars}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python count_docx_words.py /path/to/folder")
        sys.exit(1)

    main(sys.argv[1])